import docx

def Gen_docx_template(template_filename: str, output_filename: str, sections : list[tuple[str, str]]) -> None:
    '''
        This function simply generates a docx file following the given template. 
        It is rather basic since we are only obtaining text from videos, but this implementation allows
        for further customization.
        
        Parameters
        ----------------
        template_filename -> path to the template from which the summary will be generated.
        output_filename -> path where to store the generated output.
        
        sections -> In the structure of [(Title, text), (Title, text), ...]. There will be a page break between sections
        and a significant vertical gap between the title and text.
    '''
    document = docx.Document(template_filename)
    for s in sections:
        document.add_paragraph(s[0], style='Title')
        document.add_paragraph('', style='Normal')
        document.add_paragraph('', style='Normal')
        document.add_paragraph('', style='Normal')
        document.add_paragraph('', style='Normal')
        document.add_paragraph('', style='Normal')
        document.add_paragraph(s[1], style='Normal')
        document.add_page_break()
    
    document.save(output_filename)
    pass


def Gen_txt_template(output_filename: str, sections : list[tuple[str, str]]) -> None:
    '''
        This function simply generates a txt file containing the output from the information extraction.
        
        Parameters
        ----------------
        output_filename -> path where to store the generated output.
        
        sections -> In the structure of [(Title, text), (Title, text), ...].
    '''
    f = open(output_filename, 'w')
    for s in sections:
        f.write(s[0] + '\n-------------------------\n')
        f.write(s[1] + '\n::::::::::::::::::::::::::::::::\n\n\n')
    f.close()
    pass


if __name__ == '__main__':
    sections = [('SUMMARY TITLE', 'lorem ipsum...'), ('CONTENT TITLE 1', 'lorem ipsum...')]
    Gen_docx_template('template.docx', 'result.docx', sections=sections)
    Gen_txt_template('result.txt', sections=sections)